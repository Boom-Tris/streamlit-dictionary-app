import streamlit as st
import sqlite3
from docx import Document
from io import BytesIO
import requests

# เชื่อมต่อกับฐานข้อมูล SQLite
conn = sqlite3.connect('dictionary.db')
cursor = conn.cursor()

# สร้างตารางถ้ายังไม่มี


# ฟังก์ชันสำหรับการแสดงคำศัพท์ทั้งหมด
def show_terms(search_term="", lecture_filter=""):
    query = "SELECT * FROM terms WHERE 1=1"
    params = []
    
    if search_term:
        query += " AND word LIKE ?"
        params.append(f"%{search_term}%")
    
    if lecture_filter and lecture_filter != "ทั้งหมด":
        query += " AND lecture = ?"
        params.append(lecture_filter)
    
    query += " ORDER BY word ASC"
    cursor.execute(query, params)
    return cursor.fetchall()

# ฟังก์ชันสำหรับการเพิ่มคำศัพท์ใหม่
def add_term(word, definition, lecture="Lecture 1"):
    cursor.execute("INSERT INTO terms (word, definition, lecture) VALUES (?, ?, ?)", (word, definition, lecture))
    conn.commit()

# ฟังก์ชันสำหรับการแก้ไขคำศัพท์
def update_term(id, word, definition, lecture):
    cursor.execute("UPDATE terms SET word = ?, definition = ?, lecture = ? WHERE id = ?", (word, definition, lecture, id))
    conn.commit()

# ฟังก์ชันสำหรับการลบคำศัพท์
def delete_term(id):
    cursor.execute("DELETE FROM terms WHERE id = ?", (id,))
    conn.commit()

# ฟังก์ชันเพื่อสร้างไฟล์ .docx และส่งออก
def export_terms_to_docx(terms):
    doc = Document()
    doc.add_heading('คำศัพท์', 0)

    for term in terms:
        doc.add_paragraph(f"Lecture: {term[3]}")
        doc.add_paragraph(f"Vocabulary: {term[1]}")
        doc.add_paragraph(f"Meaning: {term[2]}")
        doc.add_paragraph("\n")

    byte_io = BytesIO()
    doc.save(byte_io)
    byte_io.seek(0)
    return byte_io

# ฟังก์ชันค้นหาความหมายจาก API พจนานุกรมภาษาอังกฤษ (Oxford Dictionary API)
def search_meaning_from_api(word):
    app_id = "eacd8799"  # กรอก Application ID ของคุณ
    app_key = "7b4b968a3bf970b446b2a470b0fdc3d8"  # กรอก Application Key ของคุณ
    url = f"https://od-api-sandbox.oxforddictionaries.com/api/v2/entries/en-us/{word.lower()}"
    
    headers = {
        'app_id': app_id,
        'app_key': app_key
    }

    try:
        response = requests.get(url, headers=headers, timeout=10)
        if response.status_code == 200:
            data = response.json()
            meanings = data["results"][0]["lexicalEntries"][0]["entries"][0]["senses"]
            if meanings:
                return meanings[0]["definitions"][0]
            return "ไม่พบความหมายจาก API"
        return f"ข้อผิดพลาดจาก API: {response.status_code} - {response.text}"
    except requests.exceptions.RequestException as e:
        return f"ไม่สามารถเชื่อมต่อกับ API ได้: {e}"

# หน้าแสดงคำศัพท์
def display_terms_page():
    st.title('คำศัพท์ทั้งหมด')

    search_term = st.text_input("ค้นหาคำศัพท์", "")
    lecture_filter = st.selectbox("กรองตาม Lecture:", ["ทั้งหมด", "Lecture 1", "Lecture 2", "Lecture 3", "Lecture 4", "Lecture 5", "Lecture 6"])
    
    terms = show_terms(search_term, lecture_filter)

    if terms:
        if st.button("ส่งออกเป็นไฟล์ .docx"):
            byte_io = export_terms_to_docx(terms)
            st.download_button(
                label="ดาวน์โหลดไฟล์ .docx",
                data=byte_io,
                file_name="terms.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        for term in terms:
            col1, col2, col3, col4, col5 = st.columns([2, 3, 2, 1, 1])
            with col1:
                st.text(term[1])  # คำศัพท์
            with col2:
                st.text(term[2])  # ความหมาย
            with col3:
                st.text(term[3])  # Lecture
            with col4:
                if st.button("แก้ไข", key=f"edit_{term[0]}"):
                    edit_term(term[0])
            with col5:
                if st.button("ลบ", key=f"delete_{term[0]}"):
                    delete_term(term[0])
                    st.success(f"คำศัพท์ '{term[1]}' ถูกลบแล้ว!")
                    st.experimental_rerun()
    else:
        st.write("ไม่พบคำศัพท์ที่ค้นหา.")

# หน้าเพิ่มคำศัพท์ใหม่
# หน้าเพิ่มคำศัพท์ใหม่
# ฟังก์ชันเพิ่มคำศัพท์ใหม่
def add_term_page():
    st.title('เพิ่มคำศัพท์ใหม่')

    # ช่องป้อนคำศัพท์
    word = st.text_input("คำศัพท์:")

    # ตัวเลือก Lecture
    lecture = st.selectbox(
        "เลือก Lecture:",
        ["Lecture 1", "Lecture 2", "Lecture 3", "Lecture 4", "Lecture 5", "Lecture 6"]
    )

    # ตัวเลือกวิธีการกรอกความหมาย
    option = st.radio("เลือกวิธีการกรอกความหมาย", ["กรอกความหมายเอง", "ใช้ความหมายจาก API"])
    
    # ตรวจสอบคำศัพท์ซ้ำ
    if word:
        cursor.execute("SELECT * FROM terms WHERE word = ?", (word,))
        existing_term = cursor.fetchone()
        if existing_term:
            st.error(f"คำศัพท์ '{word}' มีอยู่แล้วในฐานข้อมูล")
            return  # หยุดการดำเนินการเมื่อพบคำศัพท์ซ้ำ

    if option == "กรอกความหมายเอง":
        # ช่องป้อนความหมาย
        definition = st.text_area("ความหมาย:")
        # ปุ่มบันทึก
        if st.button('บันทึก'):
            if word and definition:
                add_term(word, definition, lecture)
                st.success(f"เพิ่มคำศัพท์ '{word}' สำเร็จ!")
            else:
                st.error("กรุณากรอกคำศัพท์และความหมาย.")
    
    elif option == "ใช้ความหมายจาก API":
        if word:
            # เรียกใช้ API
            definition = search_meaning_from_api(word)
            st.write(f"ความหมายจาก API: {definition}")
            # ปุ่มบันทึก
            if st.button('บันทึกคำศัพท์นี้'):
                add_term(word, definition, lecture)
                st.success(f"เพิ่มคำศัพท์ '{word}' สำเร็จ!")



# สร้างเมนูให้ผู้ใช้เลือกหน้า
def main():
    st.sidebar.title("เมนู")
    selection = st.sidebar.selectbox("เลือกหน้า", ["คำศัพท์ทั้งหมด", "เพิ่มคำศัพท์ใหม่"])
    if selection == "คำศัพท์ทั้งหมด":
        display_terms_page()
    elif selection == "เพิ่มคำศัพท์ใหม่":
        add_term_page()

# เรียกใช้งานฟังก์ชันหลัก
if __name__ == "__main__":
    main()
